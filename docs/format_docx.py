"""
                cell._tc.get_or_add_tcPr().append(shading)
Reads the completed docx as the ONLY source of truth, applies clean formatting.
                    run.font.size = BODY_SIZE
                    # Bold the first row (header)
        '  <w:top w:val="single" w:sz="4" w:space="0" w:color="999999"/>'
        '  <w:left w:val="single" w:sz="4" w:space="0" w:color="999999"/>'
        '  <w:bottom w:val="single" w:sz="4" w:space="0" w:color="999999"/>'
        '  <w:right w:val="single" w:sz="4" w:space="0" w:color="999999"/>'
        '  <w:insideH w:val="single" w:sz="4" w:space="0" w:color="999999"/>'
    # Remove existing borders if any
# ── Format all tables ──
def set_table_borders(table):
# ── Format all paragraphs (apply font to every run) ──
FONT = 'Times New Roman'
BODY_PT = Pt(12)
SMALL_PT = Pt(11)
        # Keep existing bold/italic, just fix font and size
        if para.style.name.startswith('Heading'):
            level = int(para.style.name.split()[-1]) if para.style.name[-1].isdigit() else 1
# ── Page margins: 1 inch ──
        else:
            run.font.size = BODY_SIZE
from copy import deepcopy
# ── Page margins: 1 inch ──
Reads the completed docx as the ONLY source of truth, applies clean formatting
(fonts, spacing, margins, table borders), saves as new file.
# ── Default style ──
ns = doc.styles['Normal']
ns.font.name = FONT
ns.font.size = BODY_PT
ns.paragraph_format.space_after = Pt(4)
ns.paragraph_format.space_before = Pt(2)
ns.paragraph_format.line_spacing = 1.15
from docx.oxml.ns import nsdecls
# ── List Paragraph style ──
try:
    lp = doc.styles['List Paragraph']
    lp.font.name = FONT
    lp.font.size = SMALL_PT
    lp.paragraph_format.space_after = Pt(2)
    lp.paragraph_format.space_before = Pt(1)
    lp.paragraph_format.line_spacing = 1.15
except KeyError:
    pass

# ── Body style ──
try:
    bs = doc.styles['Body']
    bs.font.name = FONT
    bs.font.size = BODY_PT
    bs.paragraph_format.line_spacing = 1.15
except KeyError:
    pass
# ── Update default style ──
SECTION_HEADERS = {
    "MILESTONES:", "RESEARCH PROGRESS:", "COURSE WORK PROGRESS:",
    "ACTIVITIES DURING CURRENT YEAR", "NOTABLE ACHEIVEMENTS",
    "DEVIATIONS FROM PLAN", "PLANS FOR THE NEXT YEAR:",
    "FUNDING HISTORY:", "SIGNATURES:", "For Department Use Only"
}

def is_section_header(text):
    t = text.strip()
    for kw in SECTION_HEADERS:
        if t.upper().startswith(kw.upper()):
            return True
    return False

def add_bottom_border(para):
    pPr = para._p.get_or_add_pPr()
    borders = parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        '  <w:bottom w:val="single" w:sz="6" w:space="1" w:color="888888"/>'
        '</w:pBdr>'
    )
    pPr.append(borders)

sub_headers = ["Research Progress Summary:", "Research Activities During",
               "Course-related Activities", "Degree-related Plans",
               "Research-related Plans:", "Course Work-Related Plans",
               "Degree-related Milestones:", "List of Publications:",
               "Review Notes:", "Reviewed by:"]

# ── Format all paragraphs ──
for i, para in enumerate(doc.paragraphs):
    txt = para.text.strip()

    # Title (first two lines) — centered
    if i <= 1:
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in para.runs:
            run.font.name = FONT
            run.font.size = Pt(14) if i == 0 else Pt(12)
            run.font.bold = (i == 0)
        if i == 0:
            para.paragraph_format.space_before = Pt(6)
            para.paragraph_format.space_after = Pt(2)
        continue

    # Section headers — bold + horizontal rule
    if is_section_header(txt):
        para.paragraph_format.space_before = Pt(16)
        para.paragraph_format.space_after = Pt(4)
        add_bottom_border(para)
        for run in para.runs:
            run.font.name = FONT
            run.font.size = Pt(13)
            run.font.bold = True
            run.font.color.rgb = RGBColor(0, 0, 0)
        continue

    # Sub-headers — bold, slight spacing
    is_sub = any(txt.startswith(sh) for sh in sub_headers)
    if is_sub:
        para.paragraph_format.space_before = Pt(10)
        para.paragraph_format.space_after = Pt(3)
        for run in para.runs:
            run.font.name = FONT
            run.font.size = BODY_PT
            run.font.bold = True
        continue

    # Signature lines — extra top spacing
    if "Signature of" in txt or "Reviewed by:" in txt:
        para.paragraph_format.space_before = Pt(20)

    # List paragraphs: 11pt; everything else: 12pt
    if para.style.name == 'List Paragraph':
        for run in para.runs:
            run.font.name = FONT
            run.font.size = SMALL_PT
    else:
        for run in para.runs:
            run.font.name = FONT
            run.font.size = BODY_PT
        hs.font.size = HEADING_SIZES.get(level, Pt(12))
# ── Format tables: borders, header shading, alternating rows ──
def set_table_borders(table, color="666666"):
        hs.paragraph_format.space_before = Pt(18 if level == 1 else 12)
        hs.paragraph_format.space_after = Pt(6)
    except KeyError:
        pass
        hs.font.bold = True
        hs.paragraph_format.space_before = Pt(18 if level == 1 else 12)
        hs.paragraph_format.space_after = Pt(6)
    except KeyError:
        pass
        f'  <w:top w:val="single" w:sz="6" w:space="0" w:color="{color}"/>'
        f'  <w:left w:val="single" w:sz="6" w:space="0" w:color="{color}"/>'
        f'  <w:bottom w:val="single" w:sz="6" w:space="0" w:color="{color}"/>'
        f'  <w:right w:val="single" w:sz="6" w:space="0" w:color="{color}"/>'
        f'  <w:insideH w:val="single" w:sz="4" w:space="0" w:color="{color}"/>'
        f'  <w:insideV w:val="single" w:sz="4" w:space="0" w:color="{color}"/>'
        if para.style.name.startswith('Heading'):
            level = int(para.style.name.split()[-1]) if para.style.name[-1].isdigit() else 1
            run.font.size = HEADING_SIZES.get(level, Pt(14))
        else:
            run.font.size = BODY_SIZE

# ── Format all tables ──
def set_table_borders(table):
            tcPr = cell._tc.get_or_add_tcPr()
            margins = parse_xml(
                f'<w:tcMar {nsdecls("w")}>'
                '  <w:top w:w="40" w:type="dxa"/>'
                '  <w:left w:w="80" w:type="dxa"/>'
                '  <w:bottom w:w="40" w:type="dxa"/>'
                '  <w:right w:w="80" w:type="dxa"/>'
                '</w:tcMar>'
            )
            tcPr.append(margins)
    tbl = table._tbl
                para.paragraph_format.space_after = Pt(1)
                para.paragraph_format.space_before = Pt(1)
    tblPr = tbl.tblPr
                    run.font.name = FONT
                    run.font.size = SMALL_PT
    # Remove existing borders if any
    for old in tblPr.findall(f'{{{tblPr.nsmap.get("w", "")}}}tblBorders'):
                        run.font.color.rgb = RGBColor(0, 0, 0)
    borders = parse_xml(
                shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="D9E2F3" w:val="clear"/>')
                tcPr.append(shading)
            elif ri % 2 == 0:
                shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="F5F5F5" w:val="clear"/>')
                tcPr.append(shading)
        '  <w:left w:val="single" w:sz="4" w:space="0" w:color="999999"/>'
        '  <w:bottom w:val="single" w:sz="4" w:space="0" w:color="999999"/>'
        '  <w:right w:val="single" w:sz="4" w:space="0" w:color="999999"/>'
        '  <w:insideH w:val="single" w:sz="4" w:space="0" w:color="999999"/>'
        '  <w:insideV w:val="single" w:sz="4" w:space="0" w:color="999999"/>'
        '</w:tblBorders>'
    )
    tblPr.append(borders)

for table in doc.tables:
    set_table_borders(table)
    for ri, row in enumerate(table.rows):
        for cell in row.cells:
            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.name = FONT_NAME
                    run.font.size = BODY_SIZE
                    # Bold the first row (header)
                    if ri == 0:
                        run.font.bold = True
            # Shade header row
            if ri == 0:
                shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="E8E8E8"/>')
                cell._tc.get_or_add_tcPr().append(shading)

doc.save(OUTPUT)
print(f"Formatted document saved to:\n  {OUTPUT}")
